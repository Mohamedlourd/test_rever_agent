"""Generate report_general.docx and report_conclusions.docx for REVER analysis."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

IMAGES = os.path.join(os.path.dirname(__file__), "images")
OUT = os.path.dirname(__file__)


# ─── helpers ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_img(doc, filename, width=Inches(5.8)):
    path = os.path.join(IMAGES, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[chart: {filename}]")
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)


def add_table(doc, headers, rows, hdr_bg="1E3A5F"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        _set_cell_bg(c, hdr_bg)
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = "F0F4FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = str(val)
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9)
            _set_cell_bg(c, bg)
    doc.add_paragraph()
    return tbl


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x37, 0x51, 0x8C)


def body(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(10)
    return p


def italic_note(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def bold_para(doc, label, rest=""):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10)
    if rest:
        r2 = p.add_run(" " + rest)
        r2.font.size = Pt(10)


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(10)


def sep(doc):
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  REPORT GENERAL
# ═══════════════════════════════════════════════════════════════════════════════

def build_general():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("REVER — Logistics Cost & Revenue Analysis 2025", level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    doc.add_paragraph("Scope: full year 2025   |   Updated: May 2026   |   Prepared by: Mohamed Lourdjane Aouni")
    doc.add_paragraph()

    # ── What this document is ────────────────────────────────────────────────
    h1(doc, "What this document is")
    body(doc,
         "This is the full walkthrough of everything I did to analyse how much REVER's logistics operation "
         "actually costs versus how much merchants are charged for it. I went through all the carrier invoices, "
         "crossed them with REVER's own billing data, and built a clear picture of where the business makes money, "
         "where it loses it, and where the data has gaps.")
    body(doc,
         "I ran the analysis across eight Jupyter notebooks, each covering a different angle. This document follows "
         "the same order, includes all the charts I produced, and explains what I found at each step.")
    italic_note(doc, "A companion document — report_conclusions.docx — covers what I interpret from all this "
                "and what I think REVER should look at next. This document is strictly about what the data says.")
    sep(doc)

    # ── SECTION 1: Planning & questions ─────────────────────────────────────
    h1(doc, "1. Analysis plan and questions I set out to answer")
    body(doc,
         "Before touching any data I mapped out every business question I wanted to answer, "
         "grouped into eight thematic blocks. These questions drove the notebook structure and "
         "determined what charts I built. Here is the full list.")
    sep(doc)

    h2(doc, "Block A — Overall margin and profitability")
    bold_para(doc, "A1.", "What is the total profit generated in 2025? "
              "Compare the sum of what was billed to merchants against the sum of carrier costs. Result = total gross margin of the business.")
    bold_para(doc, "A2.", "What is the average margin per shipment in euros? "
              "Total margin divided by number of shipments. Shows how much REVER earns per return managed.")
    italic_note(doc, "Sub-question: Does high shipment volume translate into better margin, or just more cost? "
                "If margin per shipment falls as volume rises, carriers may not be giving volume discounts or "
                "the extra cost is not being passed on to merchants.")
    bold_para(doc, "A3.", "Have the prices charged to merchants kept pace with carrier cost increases? "
              "Compare the monthly evolution of average carrier cost vs average merchant price. If costs rise faster "
              "than prices, margin is being compressed over time.")
    bold_para(doc, "A4.", "Are there shipments with zero cost or zero price that should not exist? "
              "Detect anomalies: shipments with carrier cost but merchant price = 0 (guaranteed loss), or "
              "merchant price but carrier cost = 0 (incomplete data).")
    sep(doc)

    h2(doc, "Block B — Carrier analysis")
    bold_para(doc, "B1.", "Which carrier generated the most profit? Which the least? "
              "Total and average margin by carrier. Identifies which carriers are most profitable.")
    italic_note(doc, "Sub-question: Is the carrier with the lowest margin expensive, or does it handle routes "
                "that are inherently more costly? Cross-reference with B3 and C1.")
    bold_para(doc, "B2.", "Which carrier has the lowest base logistics cost? "
              "Compare average logistic_cost per carrier. Reveals which contracts are most favourable at the base tariff level.")
    italic_note(doc, "Sub-question: Which carrier has the highest or lowest custom (additional) cost? "
                "A carrier can have a low base tariff but high extra charges. Both must be considered together for the real cost.")
    bold_para(doc, "B3.", "Are carrier prices consistent, or do they vary a lot for the same type of shipment? "
              "Measure the standard deviation of cost per shipment within each carrier, for similar shipments. "
              "High variability = unpredictable pricing = operational risk.")
    bold_para(doc, "B4.", "Is there too much dependency on one single carrier? "
              "Calculate what percentage of shipments and total cost each carrier represents. If one carrier "
              "accounts for more than 60-70%, there is concentration risk.")
    bold_para(doc, "B5.", "Is there a carrier that is clearly expensive for certain types of shipments and could be replaced? "
              "Cross carrier x country x return method x average cost. If one carrier is systematically more expensive "
              "in a specific combination where another carrier also operates, there is an optimisation opportunity.")
    bold_para(doc, "B6.", "Which carrier has the largest gap between real weight and billed weight? "
              "If billed weight is systematically higher than real weight for a carrier, that carrier is over-charging "
              "due to rounding or minimum billable weight policies.")
    sep(doc)

    h2(doc, "Block C — Merchant analysis")
    bold_para(doc, "C1.", "Which merchant generated the most profit? Which the least? "
              "Total and average margin by merchant (grouped by stripe_cust_id).")
    italic_note(doc, "Sub-question: For the merchants with the highest/lowest margin, which carriers were used? "
                "If a low-margin merchant uses an expensive carrier, that may be a carrier assignment issue that can be fixed.")
    italic_note(doc, "Sub-question: Which merchant has the heaviest shipments? And the lightest? "
                "Heavy shipments cost more for carriers. If a merchant generates very heavy shipments and their price "
                "does not reflect that cost, low margin is structurally built in.")
    bold_para(doc, "C2.", "Which merchants have the best profile (high volume AND high margin)? "
              "These are the strategic accounts: they generate many shipments and leave good margin.")
    bold_para(doc, "C3.", "Which merchants are clear losers (low volume AND low margin)? "
              "Merchants that consume operational resources without return. Candidates for repricing or renegotiation.")
    bold_para(doc, "C4.", "Are there merchants who have had no volume for months? Are they churning? "
              "Detect merchants who were active and have stopped sending shipments. A signal of platform abandonment.")
    bold_para(doc, "C5.", "Which merchant has a very high adjustment-to-revenue ratio? "
              "If a merchant accumulates many adjustments (overweight, remote areas, etc.), it may indicate that "
              "their customers send packages with characteristics that trigger carrier surcharges.")
    italic_note(doc, "Sub-question: Which carrier and merchant combination has the largest gap between real weight "
                "and billed weight? A merchant with lots of weight adjustments combined with a carrier that bills "
                "aggressively on rounding is the most costly combination.")
    sep(doc)

    h2(doc, "Block D — Carrier x Merchant cross-analysis")
    bold_para(doc, "D1.", "Which carrier + merchant combination generated the most or least margin? "
              "Cross-table: rows = merchants, columns = carriers, value = average margin per shipment.")
    bold_para(doc, "D2.", "Which carrier delivers the most weight at the lowest cost per merchant? "
              "Detect if there are carriers offering better cost efficiency per kg for specific merchants.")
    bold_para(doc, "D3.", "Which carrier + merchant combination has the highest or lowest custom additional cost? "
              "Some merchants may have a high custom_cost because they have special agreements or because more "
              "surcharges are passed on. Cross-referenced with carrier, this reveals whether the issue is "
              "the merchant, the carrier, or the combination.")
    sep(doc)

    h2(doc, "Block E — Time and geography")
    bold_para(doc, "E1.", "Which month had the most billing? Are there consistently less profitable months? "
              "Volume and margin by month. Business seasonality.")
    italic_note(doc, "Sub-question: Does volume growth come with margin growth, or just with more cost? "
                "If margin per shipment falls in high-volume months, carriers are not giving volume discounts and/or "
                "they are not being passed through to merchant pricing.")
    bold_para(doc, "E2.", "Which country had the most billing? "
              "Volume and margin by country_iso. Identifies the most important markets.")
    italic_note(doc, "Sub-question: Which country, by carrier and by merchant? A carrier can be very efficient "
                "in Spain and expensive in Italy. A merchant may operate heavily in France but with low margin.")
    sep(doc)

    h2(doc, "Block F — Return methods")
    bold_para(doc, "F1.", "Which return method generated the most profit? Which routes or return types are most or least profitable? "
              "Margin by DROP_OFF_POINT vs HOME_PICK_UP vs DELIVERY_PICKUP.")
    italic_note(doc, "Sub-question: Which carrier generated the most using each return method? "
                "HOME_PICK_UP tends to be more expensive for the carrier (requires active collection). If a carrier "
                "charges a lot for it and that is not passed on correctly to the merchant, it becomes a source of negative margin.")
    italic_note(doc, "Sub-question: Which merchant generated the most using each return method? "
                "Some merchants may have HOME_PICK_UP agreed for all their customers. If that method is loss-making, "
                "that specific merchant is the root cause.")
    italic_note(doc, "Sub-question: Which carrier + merchant + return method combination generated the most or least margin? "
                "The most granular level. Allows detection of very specific combinations that destroy margin.")
    sep(doc)

    h2(doc, "Block G — Adjustments")
    bold_para(doc, "G1.", "Which type of adjustment is the most costly in total? "
              "Ranking of adjustment types (Weight Adjustment, Peak Surcharge, Over Maximum Size...) by total amount. "
              "Shows where the money goes in unexpected costs.")
    italic_note(doc, "Sub-question: Which adjustment type is the most costly in each carrier? "
                "Each carrier has its own adjustment policies. Knowing what each carrier charges most allows better negotiation.")
    italic_note(doc, "Sub-question: Which adjustment type is the most costly in each merchant and carrier combination? "
                "If a merchant consistently generates a specific type of adjustment with a specific carrier, "
                "either the merchant can change their behaviour or it signals a carrier mismatch.")
    bold_para(doc, "G2.", "Which merchant has a very high adjustment ratio relative to their total billing? "
              "Merchants where adjustments represent more than a set percentage of what they are billed. They may be "
              "underpriced at base rate and compensated with unexpected adjustments.")
    sep(doc)

    h2(doc, "Block H — Data gaps and integrity")
    bold_para(doc, "H1.", "Are there shipments in carrier costs that do not appear in merchant billing? "
              "Orders the carrier charged REVER for that REVER never invoiced to a merchant. Pure cost with no recovery.")
    italic_note(doc, "Sub-question: How many are there and how much money do they represent, per carrier and per merchant? "
                "Quantifies exactly how much is lost per carrier due to absorbed costs. A very high gap in one carrier "
                "may indicate a specific operational problem: many unused labels, many late cancellations, etc.")
    bold_para(doc, "H2.", "Are there shipments billed to the merchant that have no associated carrier cost? "
              "The reverse case: REVER charged the merchant but there is no carrier invoice. May be missing data, "
              "a carrier not included in the dataset (BRT), or an error.")
    bold_para(doc, "H3.", "Are there duplicate shipments in any of the datasets? "
              "Verify that there are no repeated tracking_ids in REVER or carrier files. A duplicate would artificially "
              "inflate costs or revenues.")
    sep(doc)

    h2(doc, "Notebook structure planned")
    body(doc, "To answer all of the above, I structured the analysis into eight dedicated notebooks:")
    add_table(doc,
              ["Notebook", "Blocks covered", "Description"],
              [
                  ["01_preparacion.ipynb", "—", "Load and clean all datasets. Build the unified table (REVER + carriers joined by tracking_id). Run once."],
                  ["02_margen_global.ipynb", "A", "Total margin, margin per shipment, price vs cost evolution, anomalies (zero price/cost)."],
                  ["03_carriers.ipynb", "B", "Profitability by carrier, base vs additional cost, price consistency, concentration, billed weight."],
                  ["04_merchants.ipynb", "C", "Profitability by merchant, winner/loser profiling, churn, adjustment ratio."],
                  ["05_cruce_carrier_merchant.ipynb", "D", "Carrier x merchant margin cross-table, cost/kg efficiency, additional costs cross-referenced."],
                  ["06_geografia_tiempo.ipynb", "E", "Seasonality, volume vs margin evolution, country analysis."],
                  ["07_return_method.ipynb", "F", "Profitability by return method, cross-referenced with carrier and merchant."],
                  ["08_ajustes_gaps.ipynb", "G + H", "Adjustment types and cost, data gaps, duplicate tracking IDs."],
              ])
    sep(doc)

    # ── SECTION 2: Data ──────────────────────────────────────────────────────
    h1(doc, "2. The data I worked with")
    body(doc, "I started with two sets of files: what REVER bills to merchants, and what each carrier bills to REVER.")

    h2(doc, "REVER files (what gets charged to merchants)")
    bold_para(doc, "invoiced_logistics2025.csv", "— 294,694 rows covering all of 2025.")
    body(doc, "This is REVER's own billing data: one row per shipment, with the tracking ID, which merchant it belongs to, "
         "which return method was used, and how much was charged. The key financial fields are subtotal (net amount billed "
         "to merchant), logistic_cost (base tariff), and custom_cost (any extra personalised charge on top).")
    body(doc, "427 distinct merchants, 79 destination countries. Spain made up 48% of volume, Italy 20%, France 9%, Germany 6%.")
    sep(doc)

    bold_para(doc, "invoiced_logistics_adjustments2025.csv", "— 120,859 rows.")
    body(doc, "These are the surcharges and adjustments that carriers bill and that REVER passes on to merchants: "
         "weight corrections, remote area fees, peak surcharges, oversized packages, and so on. Total value: "
         "EUR 411,957 across 11 months.")
    sep(doc)

    h2(doc, "Carrier files (what REVER pays)")
    body(doc, "Five carriers, five completely different file formats:")
    add_table(doc,
              ["Carrier", "Format", "Coverage 2025"],
              [
                  ["Correos", "TXT (separator) + Excel", "Jan-Dec"],
                  ["Correos Express", "Excel (.xlsx), one file per month", "Jan-Dec"],
                  ["UPS", "Single Excel file", "Jan 2025 - Jan 2026"],
                  ["GLS", "Mix of Excel and CSV", "Apr-Dec"],
                  ["BRT", "CSV (latin-1), one per month", "Jan-Dec"],
              ])
    body(doc, "The join key between REVER's data and each carrier's data is the tracking ID — the shipment barcode. "
         "Every carrier stores it under a different column name, which is why normalising the data took significant effort.")
    sep(doc)

    # ── SECTION 3: Preparing ─────────────────────────────────────────────────
    h1(doc, "3. Preparing the data (01_preparacion.ipynb)")
    body(doc, "The first notebook loads every file, normalises the columns, and builds a single joined dataset.")

    h2(doc, "Joining logic")
    body(doc, "I built a carrier_master table combining Correos, Correos Express, UPS, and GLS — all carriers where "
         "tracking IDs could be matched directly. BRT was handled separately because its tracking format is incompatible.")
    body(doc, "The join is a left join from carrier_master to REVER. This means every carrier shipment is in the base "
         "table, and REVER revenue is attached where there is a match. Shipments without a REVER match represent "
         "costs that REVER absorbed without charging the merchant.")
    sep(doc)

    add_img(doc, "01_preparacion_cell17.png")
    add_img(doc, "01_match_rate_by_carrier.png")

    h2(doc, "What matched")
    add_table(doc,
              ["Dataset", "Rows", "Notes"],
              [
                  ["REVER invoiced", "294,694", "12 months, 427 merchants"],
                  ["REVER adjustments", "120,859", "11 months, EUR 411,957 total"],
                  ["Correos", "103,803", "EUR 2.40-78.89 per shipment"],
                  ["Correos Express", "31,929", "EUR 3.42-69.83 per shipment"],
                  ["UPS", "118,652", "EUR 0-1,011 per shipment"],
                  ["GLS", "256", "EUR 10.73-7,643 per shipment (see note)"],
                  ["BRT", "41,246", "aggregate only, avg EUR 4.86, total EUR 200,475"],
                  ["Carrier master", "254,640", "Correos + CE + UPS + GLS"],
                  ["Merged (joined)", "255,230", "244,149 matched (95.7%) / 11,081 unmatched (4.3%)"],
              ])

    add_table(doc,
              ["Carrier", "Match rate", "Matched / Total"],
              [
                  ["UPS", "97.2%", "115,898 / 119,200"],
                  ["Correos", "95.1%", "98,721 / 103,833"],
                  ["Correos Express", "92.5%", "29,530 / 31,941"],
                  ["GLS", "0.0%", "0 / 256"],
              ])

    add_img(doc, "01_monthly_volume_stacked.png")

    h2(doc, "The GLS situation")
    body(doc, "GLS has 256 shipments at an average cost of ~EUR 656 per shipment, under a service called "
         "InterciudadExpress. This is completely different from standard parcel returns. The tracking ID format "
         "does not match REVER's format at all, which is why the match rate is 0%. I checked GLS Spain's public "
         "service catalogue and this service name does not appear anywhere in it — it is likely a non-standard "
         "bilateral arrangement. I will come back to this in section 11.")

    h2(doc, "The BRT situation")
    body(doc, "BRT (Bartolini, now part of Geopost/DPD Group) is the leading domestic Italian parcel carrier. "
         "It has 200+ branches across Italy and operates through the DPD network internationally.")
    body(doc, "BRT's invoice uses an internal shipment number (NUMERO SPED.) that has no equivalent in REVER's tracking "
         "format. I checked whether their barcode field (SEGNAC.) could be used instead — the format matches what "
         "REVER stores for BRT-related adjustments — but the match rate with REVER's main invoice file was only 2.1%. "
         "The reason is structural: BRT shipments go almost entirely to Italian merchants (domestic Italy returns). "
         "This type of flow does not pass through REVER's standard cross-border invoicing pipeline in the same way.")
    add_img(doc, "01_brt_monthly_cost_volume.png")
    add_table(doc,
              ["Metric", "Value"],
              [
                  ["Total cost 2025", "EUR 200,475"],
                  ["Total shipments", "41,246"],
                  ["Avg cost per shipment", "EUR 4.86"],
                  ["Per-shipment margin", "Not calculable"],
              ])
    sep(doc)

    # ── SECTION 4: Block A ───────────────────────────────────────────────────
    h1(doc, "4. Block A — Overall margin and profitability")

    h2(doc, "A1 — Total revenue vs carrier cost")
    add_img(doc, "02_A1_revenue_cost_margin.png")
    add_table(doc,
              ["Item", "Amount"],
              [
                  ["Total revenue billed to merchants", "EUR 2,064,446"],
                  ["Total carrier cost", "EUR 1,620,695"],
                  ["Gross margin", "EUR 443,751 (21.5%)"],
                  ["Absorbed cost (no merchant invoice)", "EUR 250,647"],
              ])
    body(doc, "The 21.5% gross margin is the headline number. But it only covers matched shipments. The EUR 250,647 "
         "of absorbed cost is real money paid to carriers without any corresponding merchant invoice. Once that is "
         "factored in, the effective margin drops to around EUR 193,000 — closer to a 9-10% margin over what was "
         "actually billed. I break this down more carefully in the A5 section.")
    italic_note(doc, "Internal cross-check: the EUR 250,647 absorbed cost here matches exactly the total of carrier-absorbed costs in section 11 (H1).")
    sep(doc)

    h2(doc, "A2 — Average margin per shipment")
    add_img(doc, "02_A2_volume_vs_avg_margin.png")
    body(doc, "Average margin per shipment: EUR 1.82   |   Median margin %: 16.4%")
    body(doc, "The chart shows volume per month as bars and average margin per shipment as a line. It lets me see "
         "whether high-volume months compress the per-shipment margin.")
    sep(doc)

    h2(doc, "A3 — Price vs cost evolution")
    add_img(doc, "02_A3_price_vs_cost_evolution.png")
    body(doc, "Monthly average: what was charged to merchants per shipment vs what carriers charged. The area between "
         "the two lines is the margin band. When that band narrows, margin is being compressed. When it widens, the "
         "business is doing better.")
    sep(doc)

    h2(doc, "A4 — Anomalous shipments")
    add_img(doc, "02_A4_anomalies.png")
    add_table(doc,
              ["Type", "Count", "Financial impact"],
              [
                  ["Carrier cost = 0, revenue present", "0", "EUR 0"],
                  ["Revenue = 0, carrier cost present", "33", "EUR 2,026 unrecovered cost"],
                  ["Margin < 0 (loss per shipment)", "56,035", "-EUR 240,578"],
              ])
    body(doc, "56,035 shipments — 22.9% of all matched shipments — where the merchant was charged less than the carrier "
         "charged REVER. This is a EUR 240,578 direct loss spread across nearly a quarter of all shipments. "
         "This is the single most important number in the whole analysis.")
    sep(doc)

    h2(doc, "A5 — True effective margin")
    add_img(doc, "02_A5_true_effective_margin.png")
    add_table(doc,
              ["Step", "Amount", "Margin %"],
              [
                  ["Gross margin (matched shipments)", "EUR 443,751", "21.5%"],
                  ["Less: absorbed carrier cost (H1)", "-EUR 250,647", "—"],
                  ["Effective margin (ex-absorbed)", "EUR 193,104", "~9.4%"],
                  ["+ BRT estimated net (revenue EUR 319,797 - cost EUR 200,475)", "+EUR 119,322", "—"],
                  ["Conservative total net margin", "EUR 312,426", "~13-14%"],
              ])
    body(doc, "The BRT estimate is exactly that — an estimate. I know what REVER billed for those 52,081 shipments "
         "(EUR 319,797 from H2) and what BRT charged (EUR 200,475), but I cannot confirm they are the same set "
         "of shipments. I flag this as uncertain throughout.")
    sep(doc)

    h2(doc, "A6 — Where the negative-margin shipments come from")
    add_img(doc, "02_A6_negative_margin_deepdive.png")
    body(doc, "Breaking down the 56,035 loss-making shipments by carrier, by month, and by merchant. "
         "The monthly breakdown shows whether losses spike in Q4, which would be consistent with UPS peak season "
         "surcharges kicking in from late September (an additional EUR 0.20-475 per package applied Sep 29 - Jan 17). "
         "The top 10 merchants by total loss show whether this is concentrated or spread across the portfolio.")
    sep(doc)

    # ── SECTION 5: Block B ───────────────────────────────────────────────────
    h1(doc, "5. Block B — Carrier analysis")

    h2(doc, "B1 — Margin by carrier")
    add_img(doc, "03_B1_margin_by_carrier.png")
    add_table(doc,
              ["Carrier", "Shipments", "Total margin", "Avg margin/shipment", "Margin %"],
              [
                  ["UPS", "115,898", "EUR 431,299", "EUR 3.72", "26.8%"],
                  ["Correos", "98,721", "EUR 14,276", "EUR 0.14", "4.2%"],
                  ["Correos Express", "29,530", "-EUR 1,823", "-EUR 0.06", "-1.5%"],
              ])
    italic_note(doc, "Cross-check: 431,299 + 14,276 - 1,823 = EUR 443,752 vs A1 gross margin of EUR 443,751 (EUR 1 rounding)")
    body(doc, "UPS is the engine of profitability. Correos generates a thin but positive margin. "
         "Correos Express is net loss-making across its entire volume in 2025.")
    sep(doc)

    h2(doc, "B2 — Base cost vs custom cost by carrier")
    add_img(doc, "03_B2_base_vs_custom_cost.png")
    body(doc, "For each carrier: average carrier invoice cost vs the logistic_cost and custom_cost components in "
         "REVER's own pricing. This shows whether the base tariff and custom surcharge structure tracks what carriers "
         "actually charge, or whether there are systematic gaps.")
    sep(doc)

    h2(doc, "B3 — Cost variability by carrier")
    add_img(doc, "03_B3_cost_variability_boxplot.png")
    body(doc, "Box plot of carrier cost per shipment, capped at the 99th percentile so extreme outliers do not dominate. "
         "UPS shows the widest spread — expected, given the range of package types, weights, destinations, and surcharges. "
         "High variability in UPS means that average cost per UPS shipment can mislead: some shipments cost EUR 3 and some cost EUR 150+.")
    sep(doc)

    h2(doc, "B4 — Carrier concentration")
    add_img(doc, "03_B4_concentration_pie.png")
    body(doc, "UPS: 47.4% of shipments, 93.9% of gross margin\n"
         "Correos: 40.4% of shipments, 3.1% of margin\n"
         "Correos Express: 12.1% of shipments, -0.4% of margin (negative)")
    body(doc, "Almost all of the margin comes from one carrier relationship. Correos is essential for volume but "
         "contributes almost nothing to profit. CE actively reduces profitability.")
    sep(doc)

    h2(doc, "B5 — Carrier cost by destination country")
    add_img(doc, "03_B5_cost_by_country.png")
    body(doc, "Average carrier cost per shipment broken down by top 10 destination countries, grouped by carrier. "
         "This shows where a carrier is systematically more expensive than alternatives for the same route, and "
         "where substitution might make sense.")
    sep(doc)

    h2(doc, "B6 — Billed weight by carrier")
    add_img(doc, "03_B6_billed_weight.png")
    body(doc, "Only CE and UPS have billed weight data. UPS tends to bill higher weights than actual, which is "
         "explained by their published volumetric weight policy: for international shipments, UPS uses a divisor of "
         "5,000 (L cm x W cm x H cm / 5,000 = volumetric kg) and bills whichever is higher — real weight or volumetric weight. "
         "Since August 2025, UPS also rounds each dimension up to the nearest whole centimetre before calculating, "
         "which can add 5-20% to the billed weight on packages with fractional measurements. This is standard UPS policy, not a billing error.")
    sep(doc)

    h2(doc, "B7 — Correos Express: loss analysis by route")
    add_img(doc, "03_B7_ce_loss_analysis.png")
    body(doc, "Since CE is net loss-making, I looked at whether the losses are spread evenly or concentrated in "
         "specific routes or merchants. CE was previously called Chronoexpres before being acquired by Correos Group "
         "in 2011. It is a dedicated express/urgent service (24-48h delivery) that is structurally more expensive than "
         "standard Correos. The analysis here breaks CE margin by destination country and shows whether any merchants "
         "are exclusively locked into CE with no alternative.")
    sep(doc)

    # ── SECTION 6: Block C ───────────────────────────────────────────────────
    h1(doc, "6. Block C — Merchant analysis")

    h2(doc, "C1 — Top and bottom merchants by margin")
    add_img(doc, "04_C1_top_bottom_merchants.png")
    body(doc, "408 merchants had at least one matched shipment in 2025 (out of 427 total). The chart shows top 20 "
         "and bottom 20 by total margin. Bar colours represent each merchant's primary carrier.")
    sep(doc)

    h2(doc, "C2 & C3 — Winner / loser quadrant")
    add_img(doc, "04_C2C3_winner_loser_quadrant.png")
    body(doc, "Scatter plot: X-axis is shipment volume (log scale), Y-axis is average margin per shipment. "
         "The chart is divided into quadrants by the median values of both dimensions:")
    bullet(doc, "Top-right (winners): high volume, high margin — accounts to keep and grow")
    bullet(doc, "Bottom-right: large merchants that are not very profitable per shipment")
    bullet(doc, "Top-left: smaller merchants where each shipment is profitable")
    bullet(doc, "Bottom-left (losers): low volume and low margin")
    sep(doc)

    h2(doc, "C4 — Churn detection")
    add_img(doc, "04_C4_churn_detection.png")
    add_table(doc,
              ["Segment", "Merchants"],
              [
                  ["Active (had activity in last 2 months of 2025)", "318"],
                  ["Potentially churned", "90"],
                  ["Never active more than 1 month", "30"],
              ])
    body(doc, "90 merchants — 22.1% of the portfolio — stopped sending shipments before the end of the year. "
         "30 of them never made it past their first month.")
    sep(doc)

    h2(doc, "C5 — Adjustment ratio per merchant")
    add_img(doc, "04_C5_adjustment_ratio.png")
    body(doc, "For each merchant: total adjustments billed as a percentage of their total base revenue. "
         "A high ratio — anything above 20% — means the merchant consistently generates extra charges that inflate "
         "the real cost. Merchants in this zone are worth reviewing: either their shipments attract surcharges or "
         "their custom pricing does not account for those surcharges correctly.")
    sep(doc)

    h2(doc, "C6 — Were the churned merchants profitable?")
    add_img(doc, "04_C6_churned_profitability.png")
    body(doc, "Not all churn is equal. If the 90 merchants who left were loss-making accounts, their departure "
         "actually improves the overall margin. If they were profitable accounts, it is a loss of real value. "
         "The histogram shows the total margin distribution of churned merchants. The bar chart compares "
         "average margin per shipment between active and churned merchants.")
    sep(doc)

    # ── SECTION 7: Block D ───────────────────────────────────────────────────
    h1(doc, "7. Block D — Carrier x Merchant")

    h2(doc, "D1 — Margin heatmap: carrier x top 30 merchants")
    add_img(doc, "05_D1_margin_heatmap.png")
    body(doc, "Heatmap with the top 30 merchants by volume on rows and carriers on columns. Each cell shows "
         "average margin per shipment for that combination. Green is positive, red is loss. Empty cells mean "
         "that merchant does not use that carrier. Useful for spotting specific combinations that are systematically unprofitable.")
    sep(doc)

    h2(doc, "D2 — Cost per kg: carrier x top 30 merchants")
    add_img(doc, "05_D2_cost_per_kg_heatmap.png")
    body(doc, "Same structure but showing average carrier cost per kg. Only CE and UPS have weight data. "
         "Helps spot merchants whose shipments are systematically heavy and whether one carrier handles "
         "that weight profile cheaper.")
    sep(doc)

    h2(doc, "D3 — Custom cost: carrier x top 30 merchants")
    add_img(doc, "05_D3_custom_cost_heatmap.png")
    body(doc, "Average rever_custom_cost per shipment by carrier and merchant. Merchants with high custom costs "
         "on a specific carrier have some kind of personalised surcharge arrangement for that combination.")
    sep(doc)

    # ── SECTION 8: Block E ───────────────────────────────────────────────────
    h1(doc, "8. Block E — Geography and seasonality")

    h2(doc, "E1 — Monthly trend: full year 2025")
    add_img(doc, "06_E1_monthly_seasonality.png")
    body(doc, "The dataset covers all 12 months of 2025 with no gaps. The dual-panel chart shows: "
         "top panel — monthly revenue, carrier cost, and the margin band between them; "
         "bottom panel — monthly shipment volume and average margin per shipment. "
         "This is where I can see whether volume peaks in Q4 coincide with margin compression from UPS peak surcharges.")
    sep(doc)

    h2(doc, "E2 — By destination country")
    add_img(doc, "06_E2_country_volume_margin.png")
    body(doc, "Volume and average margin per shipment for the top 20 destination countries. "
         "High-volume countries with good margins are the healthiest routes; high-volume countries "
         "with low or negative margins need attention.")
    add_img(doc, "06_E2_country_carrier_heatmap.png")
    body(doc, "Heatmap of average carrier cost per shipment by country x carrier (top 10 countries). "
         "If a country is served by multiple carriers at very different costs, there may be savings available "
         "by shifting volume to the cheaper option.")
    sep(doc)

    # ── SECTION 9: Block F ───────────────────────────────────────────────────
    h1(doc, "9. Block F — Return methods")

    h2(doc, "F1 — Margin by return method")
    add_img(doc, "07_F1_margin_by_return_method.png")
    add_table(doc,
              ["Return method", "Shipments", "Avg margin/shipment", "Avg carrier cost", "Avg revenue"],
              [
                  ["DROP_OFF_POINT", "186,556 (76.4%)", "EUR 1.98", "EUR 6.33", "EUR 8.31"],
                  ["HOME_PICK_UP", "56,275 (23.0%)", "EUR 1.35", "EUR 7.63", "EUR 8.98"],
                  ["DELIVERY_PICKUP", "1,318 (0.5%)", "-EUR 0.63", "EUR 7.31", "EUR 6.67"],
              ])
    italic_note(doc, "Cross-check: 186,556 + 56,275 + 1,318 = 244,149 = total matched shipments")
    body(doc, "DROP_OFF_POINT is the best-margin method — consumers bring the package to a pickup point, which "
         "is cheaper for the carrier. DELIVERY_PICKUP is the only method with a negative average margin: "
         "the carrier costs EUR 7.31 on average but merchants are charged EUR 6.67. "
         "Small volume (1,318 shipments, -EUR 834 total) but a product priced below cost.")
    sep(doc)

    add_img(doc, "07_F1why_method_x_carrier.png")
    body(doc, "The carrier x method heatmap shows whether DELIVERY_PICKUP's deficit is consistent across all "
         "carriers or concentrated in one. If one carrier is responsible for most of the loss, that is a routing "
         "or pricing fix. If it is all carriers, the problem is the base tariff charged for this method.")
    sep(doc)

    add_img(doc, "07_F1why_method_x_merchant.png")
    body(doc, "The method x merchant heatmap (top 20 merchants) shows which merchants use DELIVERY_PICKUP most — "
         "they would be the source of most of those losses.")
    sep(doc)

    # ── SECTION 10: Block G ──────────────────────────────────────────────────
    h1(doc, "10. Block G — Adjustments")

    h2(doc, "G1 — Adjustment types by total cost")
    add_img(doc, "08_G1_adjustment_types.png")
    body(doc, "Top 20 adjustment types ranked by total amount billed across all carriers in 2025. "
         "Weight adjustments, peak surcharges, and remote area fees tend to dominate.")
    body(doc, "One specific data point worth noting: a single UPS shipment received an Over Maximum Size "
         "adjustment of EUR 567 and a Peak surcharge Over Max of EUR 532.40. This is explained by UPS's "
         "published peak season tariff: during September 29 - January 17, UPS added EUR 475.00 per package "
         "for shipments exceeding maximum limits, on top of the regular EUR 49.50 charge. Full breakdown:")
    add_table(doc,
              ["Surcharge type", "Additional amount"],
              [
                  ["Base (all packages)", "+EUR 0.20/package"],
                  ["Additional Handling (>32 kg, special dimensions)", "+EUR 7.35"],
                  ["Large Package (L+circumference >300 cm)", "+EUR 77.60"],
                  ["Over Maximum Limits (>70 kg, >274 cm, >400 cm total)", "+EUR 475.00"],
              ])
    body(doc, "These peak surcharges also explain why there are more negative-margin shipments in Q4: UPS peak "
         "costs hit from late September and the fixed merchant tariffs cannot absorb them in real time.")
    sep(doc)

    add_img(doc, "08_G1why_adj_by_carrier.png")
    body(doc, "The carrier breakdown of adjustment types is important context for contract negotiations. "
         "Each carrier has different surcharge policies, and knowing which types cost most per carrier "
         "allows targeting those clauses specifically.")
    sep(doc)

    # ── SECTION 11: Block H ──────────────────────────────────────────────────
    h1(doc, "11. Block H — Data gaps")

    h2(doc, "H1 — Absorbed costs: carrier invoices with no REVER revenue")
    add_img(doc, "08_H1_absorbed_cost.png")
    add_table(doc,
              ["Carrier", "Unmatched shipments", "Absorbed cost", "% of carrier total"],
              [
                  ["GLS", "256", "EUR 167,956", "100%"],
                  ["UPS", "3,302", "EUR 55,140", "2.77%"],
                  ["Correos", "5,112", "EUR 17,069", "4.92%"],
                  ["Correos Express", "2,411", "EUR 10,482", "7.55%"],
                  ["Total", "11,081", "EUR 250,647", "—"],
              ])
    italic_note(doc, "Cross-check: this EUR 250,647 matches exactly the absorbed cost figure in A1.")
    body(doc, "GLS's 100% unmatched rate is a special case (the non-standard B2B freight service). For standard "
         "parcel carriers, CE has the highest unmatched rate at 7.55%, meaning that for roughly 1 in 13 CE "
         "shipments REVER paid the carrier but never billed the merchant.")
    sep(doc)

    h2(doc, "H2 — REVER revenue with no carrier record")
    add_img(doc, "08_H2_revenue_no_carrier.png")
    add_table(doc,
              ["Return method", "Shipments", "Revenue"],
              [
                  ["DROP_OFF_POINT", "38,853", "EUR 242,662"],
                  ["HOME_PICK_UP", "13,069", "EUR 76,034"],
                  ["DELIVERY_PICKUP", "159", "EUR 1,101"],
                  ["Total", "52,081", "EUR 319,797"],
              ])
    body(doc, "52,081 REVER shipments have no matching carrier record. REVER invoiced merchants EUR 319,797 for these, "
         "but there is no carrier invoice to cross-reference. The most likely explanation is that most of these "
         "are BRT (Italian domestic returns) — if that is the case, the BRT cost (EUR 200,475) and this revenue "
         "(EUR 319,797) may represent the same shipments, implying a net positive margin for BRT. "
         "This cannot be confirmed without better BRT data.")
    sep(doc)

    h2(doc, "H3 — Duplicate tracking IDs")
    add_img(doc, "08_H3_duplicates.png")
    add_table(doc,
              ["Dataset", "Unique duplicated IDs", "Total duplicate rows"],
              [
                  ["REVER", "586", "4,543"],
                  ["Carrier master", "2,203", "4,501"],
              ])
    body(doc, "By carrier (duplicate rows): UPS 4,222 / GLS 174 / Correos 81 / CE 24")
    body(doc, "In the carrier data, UPS duplicates are expected: UPS issues multiple invoice lines for the same "
         "shipment (base charge, then an adjustment, then a credit). In REVER's data, 586 tracking IDs appear "
         "more than once. Some may be legitimate (same shipment billed in two different monthly periods), "
         "but the ones that appear in multiple periods deserve a closer look.")
    sep(doc)

    h2(doc, "H4 — GLS cost detail")
    add_img(doc, "08_H4_gls_detail.png")
    body(doc, "The 256 GLS shipments cost between EUR 10.73 and EUR 7,643 per shipment, with a mean of ~EUR 656. "
         "The distribution chart shows whether this is a handful of extreme outliers or a genuinely expensive service. "
         "The monthly breakdown shows in which months GLS activity occurred.")
    body(doc, "The service name InterciudadExpress does not appear anywhere in GLS Spain's public service catalogue. "
         "GLS's standard services cap out at 40 kg per parcel. This does not look like standard parcel returns.")
    sep(doc)

    h2(doc, "H5 — Duplicate tracking ID audit")
    add_img(doc, "08_H5_duplicate_billing.png")
    body(doc, "Breaking down the 586 REVER duplicate tracking IDs into two categories:")
    bullet(doc, "Multi-period duplicates: same tracking ID appears in different billing months")
    bullet(doc, "Same-period duplicates: multiple rows in the same month — likely a processing or import error")
    body(doc, "The chart shows the excess revenue associated with each type, and the distribution of how many times "
         "each duplicated ID appears.")
    sep(doc)

    # ── SECTION 12: Additional questions ─────────────────────────────────────
    h1(doc, "12. Questions that came up after reviewing everything")
    body(doc, "After going through all eight blocks, several things did not fully add up and needed more investigation. "
         "I built additional cells for each of these in the relevant notebooks.")
    sep(doc)

    h2(doc, "X1 — The effective margin is lower than 21.5%")
    add_img(doc, "02_A5_true_effective_margin.png")
    body(doc, "The headline gross margin of 21.5% is computed only on matched shipments. It does not include the "
         "EUR 250,647 of absorbed carrier costs or give visibility on BRT. When I layered everything in (cell A5), "
         "the effective margin on the matched side drops to ~9.4%. With a conservative BRT estimate added, it rises "
         "to something in the 13-14% range.")
    sep(doc)

    h2(doc, "X2 — 22.9% negative-margin shipments: where exactly?")
    add_img(doc, "02_A6_negative_margin_deepdive.png")
    body(doc, "Cell A6 breaks down the 56,035 loss-making shipments by carrier, month, and merchant. This lets me "
         "see whether the losses are random or systematic. If they cluster in Q4 (UPS peak surcharges) or in "
         "specific merchants, that changes what can be done about it.")
    sep(doc)

    h2(doc, "X3 — Correos Express: necessary or habit?")
    add_img(doc, "03_B7_ce_loss_analysis.png")
    body(doc, "CE is structurally a premium urgent service (ex-Chronoexpres, acquired by Correos Group in 2011, "
         "delivering in 24-48h). For standard e-commerce returns, that urgency is typically not needed. Cell B7 "
         "looks at which routes and merchants rely exclusively on CE — if they have Correos or UPS available as "
         "alternatives, there is a potential for rerouting.")
    sep(doc)

    h2(doc, "X4 — GLS: EUR 167,956 absorbed with no merchant invoice")
    add_img(doc, "08_H4_gls_detail.png")
    body(doc, "256 shipments at EUR 656 average is not standard returns logistics. I could not find InterciudadExpress "
         "in any GLS Spain public documentation. Cell H4 shows the cost distribution and monthly pattern. "
         "Whether these EUR 167,956 correspond to a revenue stream elsewhere is something that cannot be determined "
         "from this dataset alone.")
    sep(doc)

    h2(doc, "X5 — Were the churned merchants profitable?")
    add_img(doc, "04_C6_churned_profitability.png")
    body(doc, "90 merchants left in 2025. Cell C6 compares the margin profile of those merchants to the active ones. "
         "If churned merchants had below-average margins, their departure is neutral or positive. If they were "
         "above-average, it represents a real commercial loss worth investigating.")
    sep(doc)

    h2(doc, "X6 — 586 REVER duplicate tracking IDs")
    add_img(doc, "08_H5_duplicate_billing.png")
    body(doc, "A tracking ID identifies a unique physical shipment. If the same ID appears more than once in "
         "REVER's billing, two scenarios: multi-period (same shipment billed in two different months) or "
         "same-period (two rows in the same month — internal processing error). Cell H5 separates "
         "these and quantifies the excess revenue associated with each type.")
    sep(doc)

    h2(doc, "X7 — UPS concentration: 93.9% of gross margin")
    body(doc, "If UPS raises tariffs by 10%, gross margin falls from EUR 443k to roughly EUR 398k. A 20% rise takes "
         "it to ~EUR 352k. And UPS already applies peak surcharges seasonally (visible in Block G). The strategic "
         "question is whether there is a viable alternative that could be developed.")
    sep(doc)

    h2(doc, "X8 — BRT: EUR 319k revenue + EUR 200k cost with no per-shipment reconciliation")
    body(doc, "I have REVER revenue without carrier match (EUR 319,797) and BRT cost without REVER match (EUR 200,475). "
         "If they are the same shipments, implied BRT margin is ~37% — the best of all carriers. If they are not, "
         "there are unexplained floating costs or revenues. Without a reconcilable tracking ID, I cannot confirm either way.")
    sep(doc)

    # ── SECTION 13: Cross-checks ─────────────────────────────────────────────
    h1(doc, "13. Data quality and internal cross-checks")
    add_table(doc,
              ["Check", "Result"],
              [
                  ["A1 gross margin = B1 sum (431,299 + 14,276 - 1,823)", "EUR 443,751 vs EUR 443,752 (EUR 1 rounding)"],
                  ["A1 absorbed cost = H1 total (167,956 + 55,140 + 17,069 + 10,482)", "EUR 250,647 = EUR 250,647"],
                  ["F1 shipment count = total matched rows", "186,556 + 56,275 + 1,318 = 244,149"],
                  ["Carrier master rows = Correos + CE + UPS + GLS", "254,640 = 103,803 + 31,929 + 118,652 + 256"],
              ])
    body(doc, "All major totals cross-check correctly. The analysis is internally consistent.")
    italic_note(doc, "For conclusions, interpretations, identified problems, and suggested next steps — see report_conclusions.docx.")
    sep(doc)

    out_path = os.path.join(OUT, "report_general.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


# ═══════════════════════════════════════════════════════════════════════════════
#  REPORT CONCLUSIONS
# ═══════════════════════════════════════════════════════════════════════════════

def build_conclusions():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("REVER — Logistics Analysis 2025: Conclusions", level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    doc.add_paragraph("Scope: full year 2025   |   Updated: May 2026   |   Prepared by: Mohamed Lourdjane Aouni")
    doc.add_paragraph()

    h1(doc, "What this document is")
    body(doc,
         "This document summarises what I found after analysing REVER's logistics costs and revenues for 2025. "
         "Every statement here comes directly from a number in the dataset. Where the data does not let me confirm "
         "something, I say so explicitly. For the full analysis — charts, methodology, and all intermediate numbers "
         "— see report_general.docx.")
    sep(doc)

    # ── Section 1: What the data confirms ────────────────────────────────────
    h1(doc, "1. What the data confirms")

    h2(doc, "The gross margin is 21.5%, but the effective margin is closer to 9-10%")
    add_img(doc, "02_A5_true_effective_margin.png")
    italic_note(doc, "See section 4 (A1 and A5) of report_general.docx")
    body(doc,
         "On the 244,149 shipments where I could match carrier cost to merchant revenue, the gross margin is "
         "EUR 443,751 (21.5%). However, EUR 250,647 in carrier costs have no corresponding merchant invoice — "
         "the carrier was paid but nothing was charged back. Once those absorbed costs are subtracted, the effective "
         "margin on the matched side drops to EUR 193,104, or roughly 9.4%.")
    body(doc,
         "There is also EUR 319,797 of REVER revenue and EUR 200,475 of BRT carrier cost that I could not "
         "cross-reference per shipment due to incompatible tracking ID formats. If those correspond to the same "
         "shipments, BRT would add approximately EUR 119,000 of net margin. I cannot confirm this.")
    sep(doc)

    h2(doc, "93.9% of the gross margin comes from UPS")
    add_img(doc, "03_B4_concentration_pie.png")
    italic_note(doc, "See section 5 (B1, B4) of report_general.docx")
    body(doc,
         "UPS generates EUR 431,299 of margin at a 26.8% rate. Correos contributes EUR 14,276 (4.2%). "
         "Correos Express is net negative at -EUR 1,823 (-1.5%). Every carrier other than UPS combined "
         "contributes less than 7% of the gross margin. This is not a problem in itself, but it is a "
         "concentration fact that affects any scenario where UPS pricing changes.")
    sep(doc)

    h2(doc, "22.9% of shipments lose money — a total of -EUR 240,578")
    add_img(doc, "02_A6_negative_margin_deepdive.png")
    italic_note(doc, "See section 4 (A4, A6) of report_general.docx")
    body(doc,
         "56,035 shipments have a negative margin. This is not random noise: it is 1 in 4 shipments where "
         "the carrier cost exceeds the merchant price. The A6 chart breaks this down by carrier, by month, "
         "and by merchant — which is where the answer to why this is happening will be.")
    sep(doc)

    h2(doc, "Correos Express has a negative net margin (-1.5%, -EUR 1,823)")
    add_img(doc, "03_B7_ce_loss_analysis.png")
    italic_note(doc, "See section 5 (B1, B7) of report_general.docx")
    body(doc,
         "Across its 29,530 matched shipments, Correos Express costs more on average than what merchants are "
         "charged. CE is a structurally more expensive service than standard Correos — it is designed for "
         "urgent 24-48h delivery. The B7 chart shows where these losses are concentrated by country and "
         "which merchants are exclusively dependent on CE.")
    sep(doc)

    h2(doc, "DELIVERY_PICKUP is priced below carrier cost (-EUR 0.63/shipment)")
    add_img(doc, "07_F1_margin_by_return_method.png")
    italic_note(doc, "See section 9 (F1) of report_general.docx")
    body(doc,
         "On average, the carrier charges EUR 7.31 per DELIVERY_PICKUP shipment but the merchant is charged "
         "EUR 6.67. Across 1,318 shipments the total loss is -EUR 834. The amount is small but the direction "
         "is consistent: this method loses money on every shipment at current prices.")
    sep(doc)

    h2(doc, "EUR 250,647 was paid to carriers with no corresponding merchant invoice")
    add_img(doc, "08_H1_absorbed_cost.png")
    italic_note(doc, "See section 11 (H1) of report_general.docx")
    body(doc,
         "11,081 shipments appear in carrier invoices but not in REVER's billing data. Breakdown:")
    add_table(doc,
              ["Carrier", "Absorbed cost", "% of that carrier"],
              [
                  ["GLS", "EUR 167,956", "100%"],
                  ["UPS", "EUR 55,140", "2.77%"],
                  ["Correos", "EUR 17,069", "4.92%"],
                  ["Correos Express", "EUR 10,482", "7.55%"],
              ])
    body(doc,
         "GLS is the largest single item (EUR 167,956) and the most unusual: all 256 GLS shipments are "
         "unmatched, the service name (InterciudadExpress) does not appear in GLS Spain's public catalogue, "
         "and the average cost per shipment is EUR 656 — far above standard parcel rates. Whether this "
         "cost has a corresponding revenue stream somewhere else in the business is not visible in this dataset.")
    sep(doc)

    h2(doc, "90 merchants (22.1%) appear to have churned")
    add_img(doc, "04_C6_churned_profitability.png")
    italic_note(doc, "See section 6 (C4, C6) of report_general.docx")
    body(doc,
         "90 merchants who were active at some point in 2025 did not appear in the last two months of data. "
         "The C6 chart compares their margin profile to active merchants. Whether their departure is neutral "
         "or negative depends on whether they were above or below average in profitability.")
    sep(doc)

    h2(doc, "586 tracking IDs appear more than once in REVER billing")
    add_img(doc, "08_H5_duplicate_billing.png")
    italic_note(doc, "See section 11 (H3, H5) of report_general.docx")
    body(doc,
         "586 tracking IDs are duplicated in REVER's invoicing data. The most relevant ones are the multi-period "
         "duplicates — the same shipment appearing in two different billing months. The H5 chart quantifies the "
         "excess revenue associated with each type. Whether those duplicates represent actual double billing or "
         "a legitimate correction would need to be verified case by case.")
    sep(doc)

    h2(doc, "BRT: EUR 200,475 in cost with no per-shipment visibility")
    add_img(doc, "08_H2_revenue_no_carrier.png")
    italic_note(doc, "See sections 3 and 4 (A5) of report_general.docx")
    body(doc,
         "BRT processed 41,246 shipments at a total cost of EUR 200,475. Because BRT uses a different "
         "tracking ID format (SEGNAC), I could not match their invoices to REVER's billing data per shipment. "
         "I cannot calculate the per-shipment margin, confirm whether the implied ~37% margin estimate is real, "
         "or audit whether the invoiced amounts are correct.")
    sep(doc)

    # ── Section 2: Problems table ─────────────────────────────────────────────
    h1(doc, "2. Things that need attention, by priority")

    h2(doc, "Highest financial impact")
    add_table(doc,
              ["Issue", "What the data shows", "Where to look"],
              [
                  ["22.9% of shipments have negative margin",
                   "-EUR 240,578 total loss",
                   "Section 4 (A4, A6)"],
                  ["EUR 250,647 absorbed — carrier paid, merchant not charged",
                   "11,081 unmatched shipments",
                   "Section 11 (H1)"],
                  ["Correos Express is net negative",
                   "-EUR 1,823 total, -1.5% margin rate",
                   "Section 5 (B1, B7)"],
                  ["DELIVERY_PICKUP priced below carrier cost",
                   "-EUR 0.63/shipment, -EUR 834 total",
                   "Section 9 (F1)"],
              ])

    h2(doc, "Structural visibility gaps")
    add_table(doc,
              ["Issue", "What the data shows", "Where to look"],
              [
                  ["93.9% of gross margin in one carrier",
                   "UPS: EUR 431,299 / Correos + CE: EUR 12,453 combined",
                   "Section 5 (B4)"],
                  ["GLS: EUR 167,956 absorbed, service not in public catalogue",
                   "100% unmatched, EUR 656/shipment avg",
                   "Section 11 (H1, H4)"],
                  ["BRT: EUR 200k cost, no per-shipment reconciliation",
                   "Tracking ID format incompatible",
                   "Section 3"],
              ])

    h2(doc, "Data integrity")
    add_table(doc,
              ["Issue", "What the data shows", "Where to look"],
              [
                  ["586 duplicate tracking IDs in REVER billing",
                   "Multi-period duplicates carry double-billing risk",
                   "Section 11 (H3, H5)"],
                  ["90 merchants stopped sending shipments",
                   "22.1% of portfolio, profitability unknown",
                   "Section 6 (C4, C6)"],
              ])
    sep(doc)

    # ── Section 3: Suggestions ────────────────────────────────────────────────
    h1(doc, "3. Suggested next steps")
    body(doc,
         "These are things I would look at based on what the data shows. They are not definitive — "
         "some of them require information about how the business operates that is not in the dataset.")
    sep(doc)

    h2(doc, "1. Understand why 56,035 shipments have negative margin")
    italic_note(doc, "Section 4 (A4, A6)")
    body(doc,
         "The A6 chart already breaks this down by carrier, month, and merchant. The first step is to read "
         "that chart and identify whether the losses are concentrated in a specific carrier, a specific time "
         "period, or a specific group of merchants. That will point to whether the root cause is a tariff "
         "that needs updating, a seasonal cost that is not being passed on, or a specific merchant relationship "
         "that is structurally unprofitable.")
    sep(doc)

    h2(doc, "2. Verify the pricing of DELIVERY_PICKUP against current carrier rates")
    italic_note(doc, "Section 9 (F1)")
    body(doc,
         "The data shows that for this return method the average carrier cost (EUR 7.31) exceeds the average "
         "merchant price (EUR 6.67). Whether the current tariff was set correctly at some point and carrier "
         "costs have since moved, or whether there is a pricing configuration issue, I cannot tell from this "
         "data alone. The straightforward check is to compare the current published tariff for this method "
         "against the actual carrier invoices.")
    sep(doc)

    h2(doc, "3. Clarify the GLS cost — whether it has a revenue counterpart somewhere")
    italic_note(doc, "Section 11 (H1, H4)")
    body(doc,
         "EUR 167,956 of carrier cost with zero revenue match is a large unresolved item. The data does not "
         "tell me whether this cost is covered by a revenue stream outside of this dataset (a different "
         "billing system, a different department) or whether it is truly unrecovered. That is a question for "
         "whoever manages the GLS relationship. I have flagged it because it is the single largest "
         "absorbed cost in the analysis.")
    sep(doc)

    h2(doc, "4. Check the DELIVERY_PICKUP tariff for the current period")
    italic_note(doc, "Section 9 (F1)")
    body(doc,
         "This is the smallest financial item (-EUR 834 total) but the clearest case: the method loses "
         "money on every shipment at current pricing. Whether to reprice it or discontinue it is a product "
         "decision, but the data makes the direction unambiguous.")
    sep(doc)

    h2(doc, "5. Review the multi-period duplicate tracking IDs")
    italic_note(doc, "Section 11 (H3, H5)")
    body(doc,
         "Among the 586 duplicate tracking IDs, the multi-period ones are the most worth looking at: "
         "they are the same shipment appearing in two different billing months. The H5 chart shows the "
         "excess revenue associated with them. Whether those represent actual double billing, a correction "
         "entry, or something else would need to be verified per case. I am flagging it because it involves "
         "real money and merchant billing.")
    sep(doc)

    h2(doc, "6. Get BRT tracking IDs in a format that can be matched to REVER data")
    italic_note(doc, "Section 3")
    body(doc,
         "Currently EUR 200k of annual carrier cost and EUR 319k of likely associated revenue cannot be "
         "reconciled per shipment. BRT already has a barcode field (SEGNAC) — the question is whether that "
         "field can be exported in a format compatible with REVER's tracking_id. Until that happens, BRT "
         "is effectively a black box in this analysis.")
    sep(doc)

    h2(doc, "7. Look at CE volume alongside the B7 chart")
    italic_note(doc, "Section 5 (B7)")
    body(doc,
         "The B7 chart shows which countries and merchants are exclusively dependent on Correos Express "
         "and have no alternative carrier. That chart is the right starting point for any decision about "
         "CE routing — I cannot assess from the data whether specific CE routes have alternatives or "
         "whether there are service or operational reasons why CE is being used.")
    sep(doc)

    h2(doc, "8. Build a simple monthly margin tracker")
    italic_note(doc, "Section 4 (A1, A2, A3)")
    body(doc,
         "Several of the issues I found — absorbed costs, negative-margin shipments, seasonal surcharges "
         "not reflected in tariffs — are not things that happened overnight. A simple monthly view of "
         "total margin, margin by carrier, percentage of negative-margin shipments, and total absorbed cost "
         "would make it possible to spot when something changes before it accumulates. The data and code "
         "to produce this already exist in the notebooks from this analysis.")
    sep(doc)

    # ── Summary table ─────────────────────────────────────────────────────────
    h1(doc, "Summary")
    add_table(doc,
              ["Finding", "Number", "Priority"],
              [
                  ["Gross margin (matched shipments)", "EUR 443,751 (21.5%)", "Reference"],
                  ["Effective margin (after absorbed costs)", "~EUR 193,104 (~9.4%)", "Reference"],
                  ["Negative-margin shipments", "56,035 (22.9%), -EUR 240,578", "High"],
                  ["Absorbed costs", "EUR 250,647", "High"],
                  ["Correos Express total loss", "-EUR 1,823 (-1.5%)", "High"],
                  ["DELIVERY_PICKUP loss", "-EUR 834 (-EUR 0.63/shipment)", "High"],
                  ["UPS margin concentration", "93.9% of gross margin", "Watch"],
                  ["GLS absorbed, service unidentified", "EUR 167,956", "Clarify"],
                  ["Churned merchants", "90 (22.1%)", "Investigate"],
                  ["BRT unreconciled", "EUR 200,475 cost / EUR 319,797 revenue", "Data gap"],
                  ["Duplicate tracking IDs", "586 (multi-period subset)", "Review"],
              ])
    sep(doc)
    italic_note(doc, "For the full analysis walkthrough — data sources, charts, methodology, and internal cross-checks — see report_general.docx.")

    out_path = os.path.join(OUT, "report_conclusions.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_general()
    build_conclusions()
    print("Done.")
